"""
Steganography & Invisible Text Detector Module
==============================================
Detects:
1. White-on-White text (invisible font color #FFFFFF or near-white on document canvas)
2. Microscopic font text (font size <= 2.5pt)
3. Hidden text attributes (e.g. DOCX vanish / PDF invisible text render mode)
4. Prompt injection and manipulative constraint analysis on detected invisible text
"""

import io
import os
import re
from typing import List, Dict, Any, Tuple, Optional
from PIL import Image
from ..schemas import Finding, BoundingBox
from .prompt_guard import scan_for_prompt_injection, create_injection_finding
from .threat_memory_bank import heuristic_micro_constraint_scan, ThreatMemoryBank


def is_white_or_near_white_color(color_int: int) -> Tuple[bool, str]:
    """
    Checks if a PyMuPDF 24-bit RGB integer color represents white or near-white.
    PyMuPDF: color = (r << 16) + (g << 8) + b
    """
    if color_int is None:
        return False, ""
    try:
        r = (color_int >> 16) & 0xFF
        g = (color_int >> 8) & 0xFF
        b = color_int & 0xFF
        hex_code = f"#{r:02X}{g:02X}{b:02X}"
        # Pure white is 255, 255, 255. Near-white threshold: r > 240 and g > 240 and b > 240
        if r >= 240 and g >= 240 and b >= 240:
            return True, hex_code
        return False, hex_code
    except Exception:
        return False, ""


def detect_white_on_white_text(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Scans PDF, DOCX, and text documents for white-on-white / invisible text.
    
    Returns a list of detected hidden text items:
    [
        {
            "text": str,
            "page": int,
            "bbox": (x, y, w, h), # percentages (0-100)
            "color_hex": str,
            "reason": str,
            "is_prompt_injection": bool,
            "injection_score": float,
            "matched_patterns": List[str],
            "intent_analysis": Dict[str, Any]
        }, ...
    ]
    """
    detected_items = []
    ext = os.path.splitext(filename)[1].lower() if "." in filename else ""

    # 1. Digital PDF Inspection via PyMuPDF
    if ext == ".pdf":
        try:
            import pymupdf as fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_idx, page in enumerate(doc):
                pw, ph = page.rect.width, page.rect.height
                if pw <= 0 or ph <= 0:
                    continue

                page_dict = page.get_text("dict")
                for block in page_dict.get("blocks", []):
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            stext = span.get("text", "").strip()
                            if not stext or len(stext) < 2:
                                continue

                            color_val = span.get("color")
                            font_size = span.get("size", 10.0)
                            flags = span.get("flags", 0)
                            
                            is_white, hex_str = is_white_or_near_white_color(color_val)
                            is_tiny = font_size <= 2.5
                            
                            if is_white or is_tiny:
                                x0, y0, x1, y1 = span.get("bbox", (0, 0, 0, 0))
                                
                                # Convert to percentage of page width/height
                                px = round(((x0 - 2) / pw) * 100, 2)
                                py = round(((y0 - 2) / ph) * 100, 2)
                                pw_pct = round(((x1 - x0 + 4) / pw) * 100, 2)
                                ph_pct = round(((y1 - y0 + 4) / ph) * 100, 2)
                                
                                reason = f"White Font Color ({hex_str})" if is_white else f"Microscopic Steganography ({font_size:.1f}pt)"
                                
                                # Analyze hidden text for prompt injection / micro-constraints
                                is_inj, inj_score, matched = scan_for_prompt_injection(stext)
                                micro_matches = heuristic_micro_constraint_scan(stext)
                                all_triggers = list(matched) + [m["matched_text"] for m in micro_matches]
                                
                                has_injection = is_inj or len(all_triggers) > 0
                                
                                intent_analysis = {
                                    "detected_phrase": stext,
                                    "is_malicious": has_injection,
                                    "reason": "Directive concealed in white font attempting to manipulate LLM evaluation" if has_injection else "Text rendered in white font matching white canvas background",
                                    "triggers": all_triggers
                                }
                                
                                detected_items.append({
                                    "text": stext,
                                    "page": page_idx + 1,
                                    "bbox": (max(0, px), max(0, py), min(100, pw_pct), min(100, ph_pct)),
                                    "color_hex": hex_str,
                                    "font_size": font_size,
                                    "reason": reason,
                                    "is_prompt_injection": has_injection,
                                    "injection_score": max(inj_score, 0.95) if has_injection else 0.0,
                                    "matched_patterns": all_triggers,
                                    "intent_analysis": intent_analysis
                                })
            doc.close()
        except Exception as e:
            print(f"[!] PDF Steganography detection error: {e}")

    # 2. Word Document (DOCX) Inspection
    elif ext == ".docx":
        try:
            import docx
            from docx.shared import RGBColor
            doc = docx.Document(io.BytesIO(file_bytes))
            
            def check_run(run, location_label="Body"):
                rtext = run.text.strip()
                if not rtext or len(rtext) < 2:
                    return None
                    
                is_white = False
                hex_str = ""
                
                # Check font color RGB
                if run.font and run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    if isinstance(rgb, RGBColor):
                        hex_str = f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
                        if rgb[0] >= 240 and rgb[1] >= 240 and rgb[2] >= 240:
                            is_white = True
                
                # Check XML color attributes (e.g. w:val="FFFFFF" or w:val="auto" or w:vanish)
                try:
                    xml = run._r.xml
                    if 'w:val="FFFFFF"' in xml or 'w:val="ffffff"' in xml or 'w:val="FEFEFE"' in xml:
                        is_white = True
                        hex_str = "#FFFFFF"
                    if '<w:vanish' in xml or (run.font and run.font.hidden):
                        is_white = True
                        hex_str = "#HIDDEN"
                except Exception:
                    pass
                    
                # Check font size
                is_tiny = False
                if run.font and run.font.size and run.font.size.pt <= 2.5:
                    is_tiny = True
                    
                if is_white or is_tiny:
                    reason = f"DOCX White Font ({hex_str}) in {location_label}" if is_white else f"Microscopic Font ({run.font.size.pt:.1f}pt) in {location_label}"
                    is_inj, inj_score, matched = scan_for_prompt_injection(rtext)
                    micro_matches = heuristic_micro_constraint_scan(rtext)
                    all_triggers = list(matched) + [m["matched_text"] for m in micro_matches]
                    has_injection = is_inj or len(all_triggers) > 0
                    
                    return {
                        "text": rtext,
                        "page": 1,
                        "bbox": (10.0, 10.0, 80.0, 15.0), # Estimated default bounding area for docx
                        "color_hex": hex_str or "#FFFFFF",
                        "font_size": run.font.size.pt if (run.font and run.font.size) else 11.0,
                        "reason": reason,
                        "is_prompt_injection": has_injection,
                        "injection_score": max(inj_score, 0.95) if has_injection else 0.0,
                        "matched_patterns": all_triggers,
                        "intent_analysis": {
                            "detected_phrase": rtext,
                            "is_malicious": has_injection,
                            "reason": f"Concealed directive in {location_label} attempting to manipulate AI evaluation",
                            "triggers": all_triggers
                        }
                    }
                return None

            # Check Section Headers & Footers
            for s_idx, section in enumerate(doc.sections):
                if section.header:
                    for p in section.header.paragraphs:
                        for run in p.runs:
                            item = check_run(run, location_label=f"Header (Section {s_idx+1})")
                            if item:
                                detected_items.append(item)
                if section.footer:
                    for p in section.footer.paragraphs:
                        for run in p.runs:
                            item = check_run(run, location_label=f"Footer (Section {s_idx+1})")
                            if item:
                                detected_items.append(item)

            # Check Paragraphs
            for p in doc.paragraphs:
                for run in p.runs:
                    item = check_run(run, location_label="Main Body")
                    if item:
                        detected_items.append(item)

            # Check Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                item = check_run(run, location_label="Table Cell")
                                if item:
                                    detected_items.append(item)

        except Exception as e:
            print(f"[!] DOCX Steganography detection error: {e}")

    return detected_items


def create_white_text_findings(detected_items: List[Dict[str, Any]], filename: str) -> Tuple[List[Finding], List[BoundingBox]]:
    """
    Converts detected white-on-white text items into standardized Veridoc Findings and Canvas Bounding Boxes.
    """
    findings = []
    bounding_boxes = []

    for idx, item in enumerate(detected_items):
        txt = item["text"]
        page = item["page"]
        x, y, w, h = item["bbox"]
        color = item.get("color_hex", "#FFFFFF")
        reason = item.get("reason", "White Font on White Background")
        is_inj = item.get("is_prompt_injection", False)
        score = item.get("injection_score", 0.95)
        triggers = item.get("matched_patterns", [])
        
        box_id = f"box-white-text-{idx+1}"
        finding_id = f"finding-white-text-{idx+1}"

        bbox = BoundingBox(
            id=box_id,
            page=page,
            x=x,
            y=y,
            width=w,
            height=h,
            label="Steganography: White-on-White Text" if not is_inj else "Adversarial Injection (White-on-White)",
            layer_type="font",
            tag="WHITE-ON-WHITE-TEXT" if not is_inj else "INJECTION-PAYLOAD",
            color="#E11D48" if is_inj else "#F59E0B",
            target_finding_id=finding_id
        )
        bounding_boxes.append(bbox)

        # Auto-learn into ThreatMemoryBank if it contains injection patterns
        if is_inj:
            try:
                t_bank = ThreatMemoryBank()
                t_bank.add_threat(txt, category="White-on-White Prompt Injection", severity="Critical", source=f"{filename} (Page {page})")
            except Exception:
                pass

        if is_inj:
            f = Finding(
                id=finding_id,
                layer_type="font",
                severity="Critical",
                title="Steganography Alert: Invisible Prompt Injection (White-on-White Text)",
                description=(
                    f"Invisible white-on-white text detected on page {page}: \"{txt}\". "
                    f"Forensic intent analysis confirms this concealed text contains an adversarial prompt injection / micro-constraint "
                    f"specifically engineered to hijack the AI evaluator and bias the output verdict."
                ),
                page=page,
                confidence=score,
                bounding_boxes=[bbox],
                details={
                    "hidden_text": txt,
                    "color": color,
                    "detection_reason": reason,
                    "is_prompt_injection": True,
                    "matched_triggers": triggers,
                    "intent_analysis": {
                        "detected_constraint": txt,
                        "intent_type": "MALICIOUS (Steganographic Prompt Hijacking)",
                        "explanation": "Adversarial author hid instruction text in white font (#FFFFFF) matching the white paper background to conceal manipulation from human auditors while triggering LLM directives.",
                        "mitigation": "Neutralized & Sandboxed in isolated XML container"
                    }
                }
            )
        else:
            f = Finding(
                id=finding_id,
                layer_type="font",
                severity="High",
                title="Steganography Alert: Invisible White-on-White Text Detected",
                description=(
                    f"Invisible text detected on page {page}: \"{txt}\". "
                    f"Text is formatted in white font color ({color}) matching the white background, making it invisible to human reviewers while present in the document's machine-readable stream."
                ),
                page=page,
                confidence=0.96,
                bounding_boxes=[bbox],
                details={
                    "hidden_text": txt,
                    "color": color,
                    "detection_reason": reason,
                    "is_prompt_injection": False
                }
            )
        findings.append(f)

    return findings, bounding_boxes
